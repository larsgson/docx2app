#!/usr/bin/env python3
"""
Build book JSON files using robust TOC-based parsing.

This version uses the lenient sequential parsing strategy from test_first_error.py
to reliably extract chapter/section structure even with numbering errors.
"""

import json
import os
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

# Configuration
INPUT_DOCX = "original-book.docx"
JSON_DIR = "chapter-viewer/public/book_content_json"
MARKDOWN_DIR = "markdown_chapters"
EXCEPTIONS_FILE = "conf/exceptions.conf"
ENABLE_MARKDOWN = True  # Enable markdown generation alongside JSON


# Import functions from test_first_error.py
def normalize_for_comparison(text):
    """Normalize text for comparison - lowercase, common OCR fixes."""
    if not text:
        return ""
    text = text.lower()
    text = re.sub(r"\b[Il](\d)", r"1\1", text)
    text = re.sub(r"\b[O](\d)", r"0\1", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def load_exceptions(exceptions_file="conf/exceptions.conf"):
    """Load TOC exceptions from configuration file."""
    exceptions = {}
    if not os.path.exists(exceptions_file):
        return exceptions
    try:
        with open(exceptions_file, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#"):
                    continue
                if "=" in line:
                    wrong_part, correct_part = line.split("=", 1)
                    wrong_part = wrong_part.strip()
                    correct_part = correct_part.strip()
                    wrong_match = re.match(r"^([\d.]+)", wrong_part)
                    correct_match = re.match(r"^([\d.]+)", correct_part)
                    if wrong_match and correct_match:
                        wrong_num = wrong_match.group(1)
                        correct_num = correct_match.group(1)
                        exceptions[wrong_num] = correct_num
    except Exception as e:
        print(f"Warning: Could not load exceptions file: {e}")
    return exceptions


def normalize_toc_text(text):
    """Normalize text for TOC parsing."""
    if not text:
        return ""
    # Remove page numbers and dots
    text = re.sub(r"\s*\.{2,}\s*\d+\s*$", "", text)
    # Clean up whitespace
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def is_toc_false_positive(text, entry_type, chapter, section=None, subsection=None):
    """Check if this is a false positive (not an actual TOC entry)."""
    # Filter out dosage patterns like "0.2 mg/kg" - must have unit immediately after number
    if re.search(r"\d+\.\d+\s*(mg|ml|kg|g(?!\w)|lb|%|cc)\b", text, re.IGNORECASE):
        return True

    # Filter out age patterns like "1.5 years"
    if re.search(r"\d+\.\d+\s*(year|month|week|day|hour)", text, re.IGNORECASE):
        return True

    # Filter out decimal numbers in context like "0.5 pour-on"
    if re.match(r"^0\.\d+\s", text):
        return True

    # Chapter 0 is likely false positives
    if chapter == 0:
        return True

    # Section titles should have substantial text after the number
    # Handle extra spaces in numbering like "3. 1" or "21. 2"
    text_after_number = re.sub(r"^\d+\.\s*\d+(?:\.\s*\d+)?\s*", "", text)
    if len(text_after_number) < 3:
        return True

    # Title should start with a capital letter or quotation mark (after number)
    # Support both regular quotes and Unicode smart quotes (U+201C, U+201D, U+2018, U+2019)
    if text_after_number and not (
        text_after_number[0].isupper()
        or text_after_number[0] in "\"''\u201c\u201d\u2018\u2019"
    ):
        return True

    return False


def extract_toc_from_document(doc):
    """Extract TOC entries directly from the document."""
    toc_entries = []
    in_toc = False
    consecutive_non_toc = 0

    for para in doc.paragraphs:
        text = para.text.strip()

        # Detect TOC section (entries with dots leading to page numbers)
        if "....." in text or "....." in text.replace(" ", ""):
            in_toc = True
            consecutive_non_toc = 0

        # Stop when we hit substantial content after TOC
        if in_toc and text and not "..." in text:
            # Check if this looks like a TOC entry without dots
            if not re.match(r"^\d+\.\d+", text):
                consecutive_non_toc += 1
                if consecutive_non_toc > 50:
                    break
                continue
            else:
                consecutive_non_toc = 0

        if in_toc and text:
            # Extract TOC entry
            normalized = normalize_toc_text(text)

            # Match chapter pattern (N.0) - handle extra spaces
            chapter_match = re.match(r"^(\d+)\.\s*0\s+(.+)", normalized)
            if chapter_match:
                chapter = int(chapter_match.group(1))
                if not is_toc_false_positive(normalized, "chapter", chapter):
                    toc_entries.append(
                        {
                            "type": "chapter",
                            "chapter": chapter,
                            "section": 0,
                            "subsection": None,
                            "title": normalized,
                        }
                    )
                continue

            # Match section pattern (N.X where X > 0) - handle extra spaces
            section_match = re.match(r"^(\d+)\.\s*(\d+)\s+(.+)", normalized)
            if section_match and int(section_match.group(2)) > 0:
                chapter = int(section_match.group(1))
                section = int(section_match.group(2))
                if not is_toc_false_positive(normalized, "section", chapter, section):
                    toc_entries.append(
                        {
                            "type": "section",
                            "chapter": chapter,
                            "section": section,
                            "subsection": None,
                            "title": normalized,
                        }
                    )
                continue

            # Match subsection pattern (N.X.Y) - handle extra spaces
            subsection_match = re.match(
                r"^(\d+)\.\s*(\d+)\.\s*(\d+)\s+(.+)", normalized
            )
            if subsection_match:
                chapter = int(subsection_match.group(1))
                section = int(subsection_match.group(2))
                subsection = int(subsection_match.group(3))
                if not is_toc_false_positive(
                    normalized, "subsection", chapter, section, subsection
                ):
                    toc_entries.append(
                        {
                            "type": "subsection",
                            "chapter": chapter,
                            "section": section,
                            "subsection": subsection,
                            "title": normalized,
                        }
                    )
                continue

    return toc_entries


def build_toc_structure(toc_entries):
    """Convert flat TOC entries into expected sequence for validation."""
    expected_sequence = []

    for entry in toc_entries:
        expected_sequence.append(
            {
                "type": entry["type"],
                "chapter": entry["chapter"],
                "section": entry.get("section", 0),
                "subsection": entry.get("subsection"),
                "title": entry["title"],
                "title_normalized": normalize_for_comparison(entry["title"]),
            }
        )

    return expected_sequence


def extract_number_and_title(text, doc, para_index):
    """Extract section number and title from text."""
    if not text:
        return None

    text = re.sub(r"\s*\.{2,}.*$", "", text).strip()

    # Handle "Chapter N.0" or "Chapter N:" patterns
    chapter_match = re.match(r"^Chapter\s+(\d+)\.0\s+(.+)", text, re.IGNORECASE)
    if chapter_match:
        text = chapter_match.group(1) + ".0 " + chapter_match.group(2)
    else:
        chapter_colon_match = re.match(r"^Chapter\s+(\d+):\s+(.+)", text, re.IGNORECASE)
        if chapter_colon_match:
            text = chapter_colon_match.group(1) + ".0 " + chapter_colon_match.group(2)

    # Check for embedded section number
    embedded_match = re.search(r"\D\.(\d+\.\d+(?:\.\d+)?)\s+([A-Z])", text)
    if embedded_match:
        start_pos = embedded_match.start(1)
        text = text[start_pos:]

    # Fix common OCR errors
    text = re.sub(r"^[Il](\d)", r"1\1", text)
    text = re.sub(r"^(\d+)\.[Il]", r"\1.1", text)
    text = re.sub(r"^(\d+)\.(\d+)\.[Il]", r"\1.\2.1", text)
    text = re.sub(r"^(\d+)\.\s+(\d+)", r"\1.\2", text)

    # Look ahead for continuation lines
    combined_text = text
    use_dotall = False

    if doc and para_index is not None:
        for offset in range(1, 6):
            next_idx = para_index + offset
            if next_idx < len(doc.paragraphs):
                next_text = doc.paragraphs[next_idx].text.strip()
                if next_text and not re.match(r"^\d+\.", next_text):
                    combined_text += " " + next_text
                else:
                    break
            else:
                break
    else:
        use_dotall = True

    # Try N.X.Y pattern (subsection)
    if use_dotall:
        match = re.match(r"^(\d+)\.(\d+)\.(\d+)\s+(.*)", combined_text, re.DOTALL)
    else:
        match = re.match(r"^(\d+)\.(\d+)\.(\d+)\s*(.*?)$", combined_text)

    if match:
        return (
            int(match.group(1)),
            int(match.group(2)),
            int(match.group(3)),
            combined_text,
        )

    # Try N.X pattern
    if use_dotall:
        match = re.match(r"^(\d+)\.(\d+)\s+(.*)", combined_text, re.DOTALL)
    else:
        match = re.match(r"^(\d+)\.(\d+)\s*(.*?)$", combined_text)

    if match:
        chapter = int(match.group(1))
        section = int(match.group(2))
        if section == 0:
            return (chapter, 0, None, combined_text)
        else:
            return (chapter, section, None, combined_text)

    return None


def extract_toc_structure(doc):
    """Extract TOC structure directly from document."""
    print("Extracting TOC from document...")
    toc_entries = extract_toc_from_document(doc)
    print(f"  Found {len(toc_entries)} TOC entries")

    expected_sequence = build_toc_structure(toc_entries)
    return expected_sequence


def find_toc_end(doc):
    """Find where TOC ends in document."""
    toc_end = 0
    in_toc = False
    consecutive_non_toc = 0

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and "....." in text:
            in_toc = True
            toc_end = i
            consecutive_non_toc = 0
        elif in_toc and text:
            consecutive_non_toc += 1
            if consecutive_non_toc > 50:
                break

    return toc_end


def get_document_elements_in_order(doc, toc_end_index):
    """Yield document elements (paragraphs, tables, and images) in document order."""
    para_map = {id(p._element): (i, p) for i, p in enumerate(doc.paragraphs)}
    table_map = {id(t._element): (i, t) for i, t in enumerate(doc.tables)}

    # Track image counter
    image_counter = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            if id(element) in para_map:
                para_index, para = para_map[id(element)]
                if para_index > toc_end_index:
                    text = para.text.strip()

                    # Check if paragraph contains images
                    if para.runs:
                        for run in para.runs:
                            # Look for drawing elements that contain images
                            for drawing in run._element.findall(
                                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                            ):
                                # Find blip (binary large image or picture)
                                blips = drawing.findall(
                                    ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                                )
                                for blip in blips:
                                    # Get the relationship ID for the image
                                    embed_key = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                                    if embed_key in blip.attrib:
                                        rId = blip.attrib[embed_key]
                                        try:
                                            image_part = doc.part.related_parts[rId]
                                            image_counter += 1
                                            yield {
                                                "type": "image",
                                                "index": image_counter,
                                                "image_part": image_part,
                                                "para_index": para_index,
                                            }
                                        except KeyError:
                                            pass  # Skip if relationship not found

                    if text:
                        yield {
                            "type": "paragraph",
                            "index": para_index,
                            "text": text,
                            "doc": doc,
                            "element": para,
                        }

        elif tag == "tbl":
            if id(element) in table_map:
                table_index, table = table_map[id(element)]

                # Check if table contains section headers in cells
                # Use more strict pattern - must start with section number pattern
                has_headers = False
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        # Must start with section pattern like "3.1 " or "24.1.6 "
                        if cell_text and re.match(r"^\d+\.\d+(?:\.\d+)?\s+", cell_text):
                            has_headers = True
                            break
                    if has_headers:
                        break

                if has_headers:
                    # Process each cell in the table for section headers
                    seen_cells = (
                        set()
                    )  # Track cell IDs to avoid duplicates from merged cells
                    for row_index, row in enumerate(table.rows):
                        for col_index, cell in enumerate(row.cells):
                            # Skip if we've already processed this cell (merged cells)
                            cell_id = id(cell._element)
                            if cell_id in seen_cells:
                                continue
                            seen_cells.add(cell_id)

                            cell_text = cell.text.strip()
                            # Must start with section pattern
                            if cell_text and re.match(
                                r"^\d+\.\d+(?:\.\d+)?\s+", cell_text
                            ):
                                # Split cell by section numbers (handles multiple entries in one cell)
                                parts = re.split(
                                    r"(?=^\d+\.\d+(?:\.\d+)?\s+)",
                                    cell_text,
                                    flags=re.MULTILINE,
                                )

                                entry_num = 0
                                for part in parts:
                                    part = part.strip()
                                    if part and re.match(
                                        r"^\d+\.\d+(?:\.\d+)?\s+", part
                                    ):
                                        entry_num += 1
                                        yield {
                                            "type": "table_cell",
                                            "index": f"T{table_index}R{row_index}C{col_index}E{entry_num}",
                                            "text": part,
                                            "doc": None,
                                            "element": cell,
                                        }
                else:
                    # Yield entire table if no headers found
                    yield {
                        "type": "table",
                        "index": table_index,
                        "element": table,
                    }


def parse_document_structure(doc, exceptions, expected_sequence=None):
    """Parse document using TOC-guided approach."""
    if expected_sequence is None:
        print("Extracting TOC structure...")
        expected_sequence = extract_toc_structure(doc)
        print(f"✓ Extracted {len(expected_sequence)} expected entries")

    print("Finding TOC end...")
    toc_end_index = find_toc_end(doc)
    print(f"✓ TOC section ends at paragraph {toc_end_index}")

    print("\nParsing document structure...")

    expected_index = 0
    found_count = 0

    # Structure: chapters[chapter_num] = {sections: {section_num: {subsections: {subsec_num: elements}}}}
    chapters = {}
    current_chapter = None
    current_section = None
    current_subsection = None

    # Track all document elements for each structural unit
    chapter_elements = {}  # chapter_num -> list of elements
    section_elements = {}  # (chapter, section) -> list of elements
    subsection_elements = {}  # (chapter, section, subsection) -> list of elements

    for source in get_document_elements_in_order(doc, toc_end_index):
        # Handle images separately - they don't have element_obj initially
        if source["type"] == "image":
            element_obj = (source["image_part"], source["index"])
        else:
            element_obj = source["element"]

        # Try to parse numbering for paragraphs and table cells
        parsed = None
        if source["type"] == "paragraph":
            text = source["text"]
            parsed = extract_number_and_title(text, source["doc"], source["index"])
        elif source["type"] == "table_cell":
            text = source["text"]
            parsed = extract_number_and_title(text, None, None)

        if parsed:
            chapter, section, subsection, full_text = parsed

            # Apply exceptions
            parsed_num = f"{chapter}.{section}"
            if subsection is not None:
                parsed_num += f".{subsection}"

            if parsed_num in exceptions:
                correct_num = exceptions[parsed_num]
                parts = correct_num.split(".")
                if len(parts) == 3:
                    chapter, section, subsection = (
                        int(parts[0]),
                        int(parts[1]),
                        int(parts[2]),
                    )
                elif len(parts) == 2:
                    chapter, section = int(parts[0]), int(parts[1])
                    subsection = None if section != 0 else None

            found_count += 1

            # For table cells with headers, create a pseudo-paragraph element
            if source["type"] == "table_cell":
                element_obj = ("table_cell", source["index"], source["text"])

            # For images, store image info
            if source["type"] == "image":
                element_obj = (source["image_part"], source["index"])

            # Determine entry type
            if section == 0:
                entry_type = "chapter"
            elif subsection is None:
                entry_type = "section"
            else:
                entry_type = "subsection"

            # Check if matches expected
            if expected_index < len(expected_sequence):
                expected = expected_sequence[expected_index]
                numbering_match = (
                    expected["chapter"] == chapter
                    and expected["section"] == section
                    and expected.get("subsection") == subsection
                )

                if numbering_match:
                    expected_index += 1
                else:
                    # Try title match
                    text_normalized = normalize_for_comparison(full_text)
                    text_title_only = re.sub(
                        r"^\d+\.\d+(?:\.\d+)?\s*", "", text_normalized
                    )

                    title_match_found = None
                    for look_idx in range(
                        expected_index, min(expected_index + 5, len(expected_sequence))
                    ):
                        look_entry = expected_sequence[look_idx]
                        look_title_normalized = look_entry["title_normalized"]
                        look_title_only = re.sub(
                            r"^\d+\.\d+(?:\.\d+)?\s*", "", look_title_normalized
                        )

                        if (
                            text_title_only in look_title_only
                            or look_title_only in text_title_only
                            or text_title_only == look_title_only
                        ) and len(text_title_only) > 3:
                            title_match_found = (look_idx, look_entry)
                            break

                    if title_match_found:
                        match_idx, match_entry = title_match_found
                        chapter = match_entry["chapter"]
                        section = match_entry["section"]
                        subsection = match_entry.get("subsection")
                        expected_index = match_idx + 1
                    else:
                        # Neither numbering nor title matches - validate it's in TOC
                        # Check if this exact numbering exists anywhere in TOC
                        found_in_toc = False
                        for check_entry in expected_sequence:
                            if (
                                check_entry["chapter"] == chapter
                                and check_entry["section"] == section
                                and check_entry.get("subsection") == subsection
                            ):
                                # Numbering exists in TOC - check if title is close enough
                                toc_title = check_entry["title_normalized"]
                                toc_title_only = re.sub(
                                    r"^\d+\.\d+(?:\.\d+)?\s*", "", toc_title
                                )

                                text_normalized = normalize_for_comparison(full_text)
                                text_title_only = re.sub(
                                    r"^\d+\.\d+(?:\.\d+)?\s*", "", text_normalized
                                )

                                if len(text_title_only) > 3 and (
                                    text_title_only in toc_title_only
                                    or toc_title_only in text_title_only
                                ):
                                    found_in_toc = True
                                break

                        if not found_in_toc:
                            # False positive - skip it
                            found_count -= 1
                            continue

            # Update current structure
            if section == 0:
                # Chapter heading
                current_chapter = chapter
                current_section = None
                current_subsection = None
                if chapter not in chapters:
                    chapters[chapter] = {"sections": {}}
                    chapter_elements[chapter] = []

                # Determine correct element type
                elem_type = (
                    "table_cell" if source["type"] == "table_cell" else "paragraph"
                )
                chapter_elements[chapter].append((elem_type, element_obj))

            elif subsection is None:
                # Section heading
                current_chapter = chapter
                current_section = section
                current_subsection = None

                if chapter not in chapters:
                    chapters[chapter] = {"sections": {}}
                    chapter_elements[chapter] = []

                if section not in chapters[chapter]["sections"]:
                    chapters[chapter]["sections"][section] = {"subsections": {}}
                    section_elements[(chapter, section)] = []

                # Determine correct element type
                elem_type = (
                    "table_cell" if source["type"] == "table_cell" else "paragraph"
                )
                section_elements[(chapter, section)].append((elem_type, element_obj))

            else:
                # Subsection heading
                current_chapter = chapter
                current_section = section
                current_subsection = subsection

                if chapter not in chapters:
                    chapters[chapter] = {"sections": {}}
                    chapter_elements[chapter] = []

                if section not in chapters[chapter]["sections"]:
                    chapters[chapter]["sections"][section] = {"subsections": {}}
                    section_elements[(chapter, section)] = []

                if (
                    subsection
                    not in chapters[chapter]["sections"][section]["subsections"]
                ):
                    chapters[chapter]["sections"][section]["subsections"][
                        subsection
                    ] = []
                    subsection_elements[(chapter, section, subsection)] = []

                # Determine correct element type
                elem_type = (
                    "table_cell" if source["type"] == "table_cell" else "paragraph"
                )
                subsection_elements[(chapter, section, subsection)].append(
                    (elem_type, element_obj)
                )
        else:
            # Non-numbered element - add to current structure
            if current_chapter is not None:
                if current_subsection is not None:
                    # Add to subsection
                    key = (current_chapter, current_section, current_subsection)
                    if key in subsection_elements:
                        subsection_elements[key].append((source["type"], element_obj))
                elif current_section is not None:
                    # Add to section
                    key = (current_chapter, current_section)
                    if key in section_elements:
                        section_elements[key].append((source["type"], element_obj))
                else:
                    # Add to chapter
                    if current_chapter in chapter_elements:
                        chapter_elements[current_chapter].append(
                            (source["type"], element_obj)
                        )

    print(f"✓ Found {found_count} numbered entries")
    print(f"✓ Organized into {len(chapters)} chapters")

    return chapters, chapter_elements, section_elements, subsection_elements


def extract_paragraph_json(para, para_index):
    """Extract paragraph data as JSON."""
    return {
        "type": "paragraph",
        "index": para_index,
        "text": para.text,
        "style": para.style.name if para.style else None,
    }


def extract_image_json(image_filename, image_path, image_index):
    """Extract image data as JSON."""
    return {
        "type": "image",
        "index": image_index,
        "filename": image_filename,
        "path": image_path,
    }


def extract_table_json(table, table_index):
    """Extract table data as JSON."""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cells.append({"text": cell.text})
        rows.append({"cells": cells})

    return {"type": "table", "index": table_index, "rows": rows}


def extract_table_cell_json(cell_data, cell_index):
    """Extract table cell header data as JSON."""
    # cell_data is a tuple: ("table_cell", index_string, text)
    if isinstance(cell_data, tuple) and len(cell_data) == 3:
        return {
            "type": "table_cell_header",
            "index": cell_data[1],
            "text": cell_data[2],
        }
    # Fallback if it's just text
    return {
        "type": "table_cell_header",
        "index": cell_index,
        "text": str(cell_data),
    }


def create_markdown_css():
    """Create CSS file for markdown styling to match chapter-viewer."""
    css_content = """/* Markdown Styling - Matches Chapter Viewer */

:root {
    --primary-color: #2c3e50;
    --secondary-color: #34495e;
    --accent-color: #3498db;
    --text-color: #2c3e50;
    --bg-color: #ffffff;
    --border-color: #e0e0e0;
    --code-bg: #f5f5f5;
}

body {
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
    line-height: 1.6;
    color: var(--text-color);
    max-width: 800px;
    margin: 0 auto;
    padding: 2rem;
    background-color: var(--bg-color);
}

h1 {
    color: var(--primary-color);
    border-bottom: 3px solid var(--accent-color);
    padding-bottom: 0.5rem;
    margin-top: 2rem;
    margin-bottom: 1rem;
    font-size: 2rem;
}

h2 {
    color: var(--secondary-color);
    margin-top: 1.5rem;
    margin-bottom: 0.75rem;
    font-size: 1.5rem;
}

h3 {
    color: var(--secondary-color);
    margin-top: 1.25rem;
    margin-bottom: 0.5rem;
    font-size: 1.25rem;
}

p {
    margin-bottom: 1rem;
    text-align: justify;
}

strong, b {
    color: var(--primary-color);
    font-weight: 600;
}

em, i {
    color: var(--secondary-color);
}

ul, ol {
    margin-bottom: 1rem;
    padding-left: 2rem;
}

li {
    margin-bottom: 0.5rem;
}

table {
    border-collapse: collapse;
    width: 100%;
    margin: 1.5rem 0;
    background-color: var(--bg-color);
    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
}

th {
    background-color: var(--primary-color);
    color: white;
    padding: 0.75rem;
    text-align: left;
    font-weight: 600;
}

td {
    padding: 0.75rem;
    border: 1px solid var(--border-color);
}

tr:nth-child(even) {
    background-color: #f9f9f9;
}

code {
    background-color: var(--code-bg);
    padding: 0.2rem 0.4rem;
    border-radius: 3px;
    font-family: "Courier New", monospace;
    font-size: 0.9em;
}

pre {
    background-color: var(--code-bg);
    padding: 1rem;
    border-radius: 5px;
    overflow-x: auto;
    margin: 1rem 0;
}

pre code {
    background-color: transparent;
    padding: 0;
}

a {
    color: var(--accent-color);
    text-decoration: none;
}

a:hover {
    text-decoration: underline;
}

.nav-links {
    display: flex;
    justify-content: space-between;
    margin: 2rem 0;
    padding: 1rem;
    background-color: #f5f5f5;
    border-radius: 5px;
}

.nav-links a {
    padding: 0.5rem 1rem;
    background-color: var(--accent-color);
    color: white;
    border-radius: 3px;
    transition: background-color 0.3s;
}

.nav-links a:hover {
    background-color: var(--primary-color);
    text-decoration: none;
}

blockquote {
    border-left: 4px solid var(--accent-color);
    margin: 1rem 0;
    padding-left: 1rem;
    color: var(--secondary-color);
    font-style: italic;
}

hr {
    border: none;
    border-top: 2px solid var(--border-color);
    margin: 2rem 0;
}

@media print {
    body {
        max-width: 100%;
        padding: 0;
    }

    .nav-links {
        display: none;
    }
}

@media (max-width: 768px) {
    body {
        padding: 1rem;
    }

    h1 {
        font-size: 1.5rem;
    }

    h2 {
        font-size: 1.25rem;
    }

    table {
        font-size: 0.9rem;
    }
}
"""
    return css_content


def create_markdown_index(chapters, markdown_dir):
    """Create README.md index file for markdown chapters."""
    lines = [
        "# Book Content - Markdown Format",
        "",
        "This directory contains the book content in Markdown format, generated from the Word document.",
        "",
        "## Chapters",
        "",
    ]

    for chapter_num in sorted(chapters.keys()):
        chapter_data = chapters[chapter_num]
        chapter_title = f"Chapter {chapter_num}"

        # Extract title from first section if available
        if chapter_data.get("sections"):
            chapter_title = f"{chapter_num}.0"

        lines.append(f"### [{chapter_title}](chapter_{chapter_num:02d}/intro.md)")
        lines.append("")

        # List sections
        for section_num in sorted(chapter_data["sections"].keys()):
            section_data = chapter_data["sections"][section_num]
            lines.append(
                f"- [{chapter_num}.{section_num}](chapter_{chapter_num:02d}/section_{section_num:02d}.md)"
            )

            # List subsections
            if section_data.get("subsections"):
                for subsection_num in sorted(section_data["subsections"].keys()):
                    lines.append(
                        f"  - [{chapter_num}.{section_num}.{subsection_num}](chapter_{chapter_num:02d}/section_{section_num:02d}_{subsection_num:02d}.md)"
                    )

        lines.append("")

    lines.extend(
        [
            "---",
            "",
            "## Viewing",
            "",
            "Open any `.md` file in a Markdown viewer or editor. For best results with styling:",
            "",
            "1. Use a Markdown viewer that supports custom CSS",
            "2. Link to `style.css` in the markdown files",
            "3. Or use the web viewer for the full interactive experience",
            "",
            "## Format",
            "",
            "- **Bold text** indicates important terms or emphasis",
            "- *Italic text* for references or subtle emphasis",
            "- Tables are formatted with markdown table syntax",
            "- Lists use standard markdown list formatting",
            "",
        ]
    )

    # Write README
    readme_path = os.path.join(markdown_dir, "README.md")
    with open(readme_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    print(f"\n✓ Created markdown index: {readme_path}")

    # Write CSS file
    css_path = os.path.join(markdown_dir, "style.css")
    with open(css_path, "w", encoding="utf-8") as f:
        f.write(create_markdown_css())

    print(f"✓ Created CSS file: {css_path}")


def format_text_markdown(text, runs=None):
    """Format text with markdown formatting based on runs."""
    if not runs:
        return text

    result = []
    for run in runs:
        run_text = run.get("text", "")
        if not run_text:
            continue

        # Apply formatting
        if run.get("bold") and run.get("italic"):
            run_text = f"***{run_text}***"
        elif run.get("bold"):
            run_text = f"**{run_text}**"
        elif run.get("italic"):
            run_text = f"*{run_text}*"

        result.append(run_text)

    return "".join(result) if result else text


def extract_paragraph_markdown(para):
    """Extract paragraph as markdown."""
    text = para.text.strip()
    if not text:
        return ""

    # Check for heading styles
    style_name = para.style.name if para.style else ""

    # Format runs for inline formatting
    runs_data = []
    for run in para.runs:
        runs_data.append(
            {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
            }
        )

    formatted_text = format_text_markdown(text, runs_data)

    # Handle list styles
    if "List Bullet" in style_name or "List" in style_name:
        return f"- {formatted_text}"
    elif "List Number" in style_name:
        return f"1. {formatted_text}"

    return formatted_text


def extract_table_markdown(table):
    """Extract table as markdown."""
    if not table.rows:
        return ""

    lines = []

    # Process rows
    for row_idx, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace("\n", " ")
            cells.append(cell_text)

        # Create table row
        lines.append("| " + " | ".join(cells) + " |")

        # Add separator after first row (header)
        if row_idx == 0:
            lines.append("| " + " | ".join(["---"] * len(cells)) + " |")

    return "\n".join(lines)


def is_wmf_image(image_data):
    """Check if image data is WMF format by checking magic bytes."""
    if len(image_data) < 4:
        return False
    # WMF magic bytes: 0xD7CDC69A or 0x01000900
    magic = image_data[:4]
    return magic == b"\xd7\xcd\xc6\x9a" or magic == b"\x01\x00\x09\x00"


def convert_wmf_to_png(wmf_path, output_path):
    """Convert WMF to PNG using LibreOffice -> PDF -> PNG chain."""
    import shutil
    import subprocess
    import tempfile
    from pathlib import Path

    # Try LibreOffice first (best for WMF)
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                # Convert WMF to PDF using LibreOffice
                result = subprocess.run(
                    [
                        soffice,
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        tmpdir,
                        wmf_path,
                    ],
                    capture_output=True,
                    text=True,
                    timeout=30,
                )

                # Find the generated PDF (LibreOffice might name it differently)
                pdf_files = list(Path(tmpdir).glob("*.pdf"))
                if pdf_files:
                    pdf_path = str(pdf_files[0])

                    # Convert PDF to PNG using ImageMagick
                    magick_cmd = None
                    if shutil.which("magick"):
                        magick_cmd = [
                            "magick",
                            "-density",
                            "300",
                            pdf_path,
                            "-flatten",
                            "png:" + output_path,
                        ]
                    elif shutil.which("convert"):
                        magick_cmd = [
                            "convert",
                            "-density",
                            "300",
                            pdf_path,
                            "-flatten",
                            "png:" + output_path,
                        ]

                    if magick_cmd:
                        result = subprocess.run(
                            magick_cmd,
                            capture_output=True,
                            text=True,
                            timeout=30,
                        )

                        if result.returncode == 0 and os.path.exists(output_path):
                            # Verify it's actually a PNG
                            with open(output_path, "rb") as f:
                                header = f.read(8)
                                if header[:4] == b"\x89PNG":
                                    return True
                            print(f"    ⚠️  Output is not PNG format")
                            return False

        except subprocess.TimeoutExpired:
            print(f"    ⚠️  WMF conversion timeout")
            return False
        except Exception as e:
            print(f"    ⚠️  LibreOffice conversion error: {e}")

    # Fallback: Try ImageMagick directly (needs WMF delegates)
    try:
        magick_cmd = None
        if shutil.which("magick"):
            magick_cmd = ["magick", wmf_path, "png:" + output_path]
        elif shutil.which("convert"):
            magick_cmd = ["convert", wmf_path, "png:" + output_path]
        else:
            print(f"    ⚠️  No conversion tools found - cannot convert WMF")
            return False

        result = subprocess.run(
            magick_cmd,
            capture_output=True,
            text=True,
            timeout=30,
        )

        if result.returncode == 0 and os.path.exists(output_path):
            # Verify it's actually a PNG
            with open(output_path, "rb") as f:
                header = f.read(8)
                if header[:4] == b"\x89PNG":
                    return True
            print(f"    ⚠️  Output is not PNG format")
            return False
        else:
            print(f"    ⚠️  WMF conversion failed: {result.stderr}")
            return False
    except subprocess.TimeoutExpired:
        print(f"    ⚠️  WMF conversion timeout")
        return False
    except Exception as e:
        print(f"    ⚠️  WMF conversion error: {e}")
        return False


def extract_and_save_image(
    image_part, chapter_num, image_index, output_dir, is_json_dir=True
):
    """Extract image and save to pictures directory. Returns (filename, path) or None."""
    # Create pictures directory
    pictures_dir = os.path.join(output_dir, f"chapter_{chapter_num:02d}", "pictures")
    os.makedirs(pictures_dir, exist_ok=True)

    # Determine file extension from content type
    content_type = image_part.content_type
    if "png" in content_type:
        ext = "png"
    elif "jpeg" in content_type or "jpg" in content_type:
        ext = "jpg"
    elif "gif" in content_type:
        ext = "gif"
    else:
        ext = "png"  # default

    # Save image
    image_filename = f"image_{image_index:03d}.{ext}"
    image_path = os.path.join(pictures_dir, image_filename)

    # Get image data
    image_data = image_part.blob

    # Check if it's WMF and needs conversion
    if is_wmf_image(image_data):
        # Save as temporary WMF file
        wmf_path = image_path + ".wmf.tmp"
        with open(wmf_path, "wb") as f:
            f.write(image_data)

        # Convert to PNG
        png_path = image_path.replace(f".{ext}", ".png")
        if convert_wmf_to_png(wmf_path, png_path):
            # Backup original WMF
            backup_path = png_path + ".wmf.backup"
            os.rename(wmf_path, backup_path)
            image_filename = os.path.basename(png_path)
            image_path = png_path
        else:
            # Conversion failed, save as-is
            os.remove(wmf_path)
            with open(image_path, "wb") as f:
                f.write(image_data)
    else:
        # Regular image, save directly
        with open(image_path, "wb") as f:
            f.write(image_data)

    # Return relative path and filename
    if is_json_dir:
        return image_filename, f"pictures/{image_filename}"
    else:
        # For markdown
        return f"pictures/{image_filename}"


def save_markdown_file(
    filepath, content_items, chapter_num, section_num=None, subsection_num=None
):
    """Save content as markdown file."""
    lines = []

    # Add HTML head with CSS link for better viewing
    lines.append('<link rel="stylesheet" href="../style.css">')
    lines.append("")

    # Add navigation breadcrumb
    nav_parts = ["[Home](../README.md)"]
    nav_parts.append(f"[Chapter {chapter_num}](intro.md)")

    if section_num is not None:
        nav_parts.append(f"Section {section_num}")
    if subsection_num is not None:
        nav_parts.append(f"Subsection {subsection_num}")

    lines.append(" → ".join(nav_parts))
    lines.append("")
    lines.append("---")
    lines.append("")

    # Add header
    if subsection_num is not None:
        lines.append(f"# {chapter_num}.{section_num}.{subsection_num}")
    elif section_num is not None:
        lines.append(f"# {chapter_num}.{section_num}")
    else:
        lines.append(f"# Chapter {chapter_num}")

    lines.append("")

    # Process content items
    for item_type, elem in content_items:
        if item_type == "paragraph":
            md_text = extract_paragraph_markdown(elem)
            if md_text:
                lines.append(md_text)
                lines.append("")
        elif item_type == "table":
            md_table = extract_table_markdown(elem)
            if md_table:
                lines.append(md_table)
                lines.append("")
        elif item_type == "table_cell":
            # Handle table cell headers
            if isinstance(elem, tuple) and len(elem) == 3:
                lines.append(f"**{elem[2]}**")
                lines.append("")
        elif item_type == "image":
            # Handle images - elem is the relative path
            if isinstance(elem, str):
                lines.append(f"![Image]({elem})")
                lines.append("")

    # Add footer navigation
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append('<div class="nav-links">')
    lines.append('<a href="../README.md">← Back to Index</a>')
    lines.append(f'<a href="intro.md">Chapter {chapter_num} Home</a>')
    lines.append("</div>")

    # Write file
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def create_navigation_index(chapters, expected_sequence, output_dir):
    """Create index.json for chapter-viewer navigation."""
    # Build a map of chapter/section numbers to titles from TOC
    title_map = {}
    for entry in expected_sequence:
        chapter = entry["chapter"]
        section = entry.get("section", 0)
        subsection = entry.get("subsection")
        title = entry["title"]

        if subsection is not None:
            key = (chapter, section, subsection)
        elif section == 0:
            key = (chapter, "chapter")
        else:
            key = (chapter, section)

        title_map[key] = title

    # Build index structure
    index_data = {"book_title": "Animal Health Handbook", "chapters": []}

    total_sections = 0

    for chapter_num in sorted(chapters.keys()):
        chapter_data = chapters[chapter_num]

        # Get chapter title
        chapter_key = (chapter_num, "chapter")
        chapter_title = title_map.get(chapter_key, f"Chapter {chapter_num}")

        chapter_obj = {"number": chapter_num, "title": chapter_title, "sections": []}

        # Add intro section (section 0)
        intro_title = chapter_title
        chapter_obj["sections"].append(
            {
                "section_number": 0,
                "title": intro_title,
                "path": f"chapter_{chapter_num:02d}/intro.json",
            }
        )

        # Add regular sections
        for section_num in sorted(chapter_data["sections"].keys()):
            section_key = (chapter_num, section_num)
            section_title = title_map.get(section_key, f"{chapter_num}.{section_num}")

            chapter_obj["sections"].append(
                {
                    "section_number": section_num,
                    "title": section_title,
                    "path": f"chapter_{chapter_num:02d}/section_{section_num:02d}.json",
                }
            )

        # Add total sections count for this chapter
        chapter_obj["total_sections"] = len(chapter_obj["sections"])
        total_sections += chapter_obj["total_sections"]

        index_data["chapters"].append(chapter_obj)

    # Add total counts
    index_data["total_chapters"] = len(index_data["chapters"])
    index_data["total_sections"] = total_sections

    # Write index.json
    index_path = os.path.join(output_dir, "index.json")
    with open(index_path, "w", encoding="utf-8") as f:
        json.dump(index_data, f, indent=2)


def build_book_json():
    """Build book JSON files and markdown from Word document."""
    print("=" * 80)
    print("BUILD BOOK - JSON and Markdown Generation")
    print("=" * 80)
    print()

    # Load exceptions
    print("Loading exceptions configuration...")
    exceptions = load_exceptions(EXCEPTIONS_FILE)
    if exceptions:
        print(f"✓ Loaded {len(exceptions)} exception(s)")
    else:
        print("✓ No exceptions configured")
    print()

    # Load document
    print(f"Loading document: {INPUT_DOCX}")
    doc = Document(INPUT_DOCX)
    print(f"✓ Loaded {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    print()

    # Extract TOC structure for navigation index
    print("Extracting TOC structure...")
    expected_sequence = extract_toc_structure(doc)
    print(f"✓ Extracted {len(expected_sequence)} expected entries")
    print()

    # Parse structure
    chapters, chapter_elements, section_elements, subsection_elements = (
        parse_document_structure(doc, exceptions, expected_sequence)
    )

    # Create output directories
    print(f"\nCreating output directory: {JSON_DIR}")
    os.makedirs(JSON_DIR, exist_ok=True)

    if ENABLE_MARKDOWN:
        print(f"Creating markdown directory: {MARKDOWN_DIR}")
        os.makedirs(MARKDOWN_DIR, exist_ok=True)

    # Track images for markdown
    image_paths = {}

    # Process each chapter
    print("\nProcessing chapters...")
    for chapter_num in sorted(chapters.keys()):
        print(f"\n  Chapter {chapter_num}:")

        chapter_dir = os.path.join(JSON_DIR, f"chapter_{chapter_num:02d}")
        os.makedirs(chapter_dir, exist_ok=True)

        if ENABLE_MARKDOWN:
            md_chapter_dir = os.path.join(MARKDOWN_DIR, f"chapter_{chapter_num:02d}")
            os.makedirs(md_chapter_dir, exist_ok=True)

        chapter_data = chapters[chapter_num]

        # Build intro section (chapter-level elements)
        intro_content = []
        para_idx = 0
        table_idx = 0

        if chapter_num in chapter_elements:
            for elem_type, elem in chapter_elements[chapter_num]:
                if elem_type == "paragraph":
                    intro_content.append(extract_paragraph_json(elem, para_idx))
                    para_idx += 1
                elif elem_type == "table":
                    intro_content.append(extract_table_json(elem, table_idx))
                    table_idx += 1
                elif elem_type == "table_cell":
                    intro_content.append(extract_table_cell_json(elem, para_idx))
                    para_idx += 1
                elif elem_type == "image":
                    # elem is a tuple: (image_part, image_index)
                    if isinstance(elem, tuple) and len(elem) == 2:
                        image_part, img_idx = elem

                        # Always save to JSON directory for chapter-viewer
                        result = extract_and_save_image(
                            image_part, chapter_num, img_idx, JSON_DIR, is_json_dir=True
                        )
                        if result:
                            image_filename, image_rel_path = result
                            intro_content.append(
                                extract_image_json(
                                    image_filename, image_rel_path, img_idx
                                )
                            )

                        # Also save to markdown directory if enabled
                        if ENABLE_MARKDOWN:
                            md_img_path = extract_and_save_image(
                                image_part,
                                chapter_num,
                                img_idx,
                                MARKDOWN_DIR,
                                is_json_dir=False,
                            )
                            if md_img_path:
                                # Store for markdown generation
                                if (chapter_num, None, None) not in image_paths:
                                    image_paths[(chapter_num, None, None)] = []
                                image_paths[(chapter_num, None, None)].append(
                                    md_img_path
                                )

        # Save intro
        if intro_content:
            intro_file = os.path.join(chapter_dir, "intro.json")
            with open(intro_file, "w", encoding="utf-8") as f:
                json.dump({"content": intro_content}, f, indent=2)
            print(f"    ✓ intro.json ({len(intro_content)} items)")

            if ENABLE_MARKDOWN and chapter_num in chapter_elements:
                md_file = os.path.join(md_chapter_dir, "intro.md")
                # Add images to content for markdown
                md_content = list(chapter_elements[chapter_num])
                key = (chapter_num, None, None)
                if key in image_paths:
                    for img_path in image_paths[key]:
                        md_content.append(("image", img_path))
                save_markdown_file(md_file, md_content, chapter_num)
                print(f"    ✓ intro.md")

        # Process sections
        for section_num in sorted(chapter_data["sections"].keys()):
            section_data = chapter_data["sections"][section_num]

            # Build section content
            section_content = []
            para_idx = 0
            table_idx = 0

            key = (chapter_num, section_num)
            if key in section_elements:
                for elem_type, elem in section_elements[key]:
                    if elem_type == "paragraph":
                        section_content.append(extract_paragraph_json(elem, para_idx))
                        para_idx += 1
                    elif elem_type == "table":
                        section_content.append(extract_table_json(elem, table_idx))
                        table_idx += 1
                    elif elem_type == "table_cell":
                        section_content.append(extract_table_cell_json(elem, para_idx))
                        para_idx += 1
                    elif elem_type == "image":
                        if isinstance(elem, tuple) and len(elem) == 2:
                            image_part, img_idx = elem

                            # Always save to JSON directory for chapter-viewer
                            result = extract_and_save_image(
                                image_part,
                                chapter_num,
                                img_idx,
                                JSON_DIR,
                                is_json_dir=True,
                            )
                            if result:
                                image_filename, image_rel_path = result
                                section_content.append(
                                    extract_image_json(
                                        image_filename, image_rel_path, img_idx
                                    )
                                )

                            # Also save to markdown directory if enabled
                            if ENABLE_MARKDOWN:
                                md_img_path = extract_and_save_image(
                                    image_part,
                                    chapter_num,
                                    img_idx,
                                    MARKDOWN_DIR,
                                    is_json_dir=False,
                                )
                                if md_img_path:
                                    if (
                                        chapter_num,
                                        section_num,
                                        None,
                                    ) not in image_paths:
                                        image_paths[
                                            (chapter_num, section_num, None)
                                        ] = []
                                    image_paths[
                                        (chapter_num, section_num, None)
                                    ].append(md_img_path)

            # Save section
            section_file = os.path.join(chapter_dir, f"section_{section_num:02d}.json")
            with open(section_file, "w", encoding="utf-8") as f:
                json.dump({"content": section_content}, f, indent=2)
            print(
                f"    ✓ section_{section_num:02d}.json ({len(section_content)} items)"
            )

            if ENABLE_MARKDOWN:
                key = (chapter_num, section_num)
                if key in section_elements:
                    md_file = os.path.join(
                        md_chapter_dir, f"section_{section_num:02d}.md"
                    )
                    # Add images to content for markdown
                    md_content = list(section_elements[key])
                    img_key = (chapter_num, section_num, None)
                    if img_key in image_paths:
                        for img_path in image_paths[img_key]:
                            md_content.append(("image", img_path))
                    save_markdown_file(md_file, md_content, chapter_num, section_num)
                    print(f"    ✓ section_{section_num:02d}.md")

            # Process subsections
            if section_data["subsections"]:
                for subsection_num in sorted(section_data["subsections"].keys()):
                    subsection_content = []
                    para_idx = 0
                    table_idx = 0

                    key = (chapter_num, section_num, subsection_num)
                    if key in subsection_elements:
                        for elem_type, elem in subsection_elements[key]:
                            if elem_type == "paragraph":
                                subsection_content.append(
                                    extract_paragraph_json(elem, para_idx)
                                )
                                para_idx += 1
                            elif elem_type == "table":
                                subsection_content.append(
                                    extract_table_json(elem, table_idx)
                                )
                                table_idx += 1
                            elif elem_type == "table_cell":
                                subsection_content.append(
                                    extract_table_cell_json(elem, para_idx)
                                )
                                para_idx += 1
                            elif elem_type == "image":
                                if isinstance(elem, tuple) and len(elem) == 2:
                                    image_part, img_idx = elem

                                    # Always save to JSON directory for chapter-viewer
                                    result = extract_and_save_image(
                                        image_part,
                                        chapter_num,
                                        img_idx,
                                        JSON_DIR,
                                        is_json_dir=True,
                                    )
                                    if result:
                                        image_filename, image_rel_path = result
                                        subsection_content.append(
                                            extract_image_json(
                                                image_filename, image_rel_path, img_idx
                                            )
                                        )

                                    # Also save to markdown directory if enabled
                                    if ENABLE_MARKDOWN:
                                        md_img_path = extract_and_save_image(
                                            image_part,
                                            chapter_num,
                                            img_idx,
                                            MARKDOWN_DIR,
                                            is_json_dir=False,
                                        )
                                        if md_img_path:
                                            if (
                                                chapter_num,
                                                section_num,
                                                subsection_num,
                                            ) not in image_paths:
                                                image_paths[
                                                    (
                                                        chapter_num,
                                                        section_num,
                                                        subsection_num,
                                                    )
                                                ] = []
                                            image_paths[
                                                (
                                                    chapter_num,
                                                    section_num,
                                                    subsection_num,
                                                )
                                            ].append(md_img_path)

                    # Save subsection
                    subsection_file = os.path.join(
                        chapter_dir,
                        f"section_{section_num:02d}_{subsection_num:02d}.json",
                    )
                    with open(subsection_file, "w", encoding="utf-8") as f:
                        json.dump({"content": subsection_content}, f, indent=2)
                    print(
                        f"      ✓ section_{section_num:02d}_{subsection_num:02d}.json ({len(subsection_content)} items)"
                    )

                    if ENABLE_MARKDOWN:
                        key = (chapter_num, section_num, subsection_num)
                        if key in subsection_elements:
                            md_file = os.path.join(
                                md_chapter_dir,
                                f"section_{section_num:02d}_{subsection_num:02d}.md",
                            )
                            # Add images to content for markdown
                            md_content = list(subsection_elements[key])
                            img_key = (chapter_num, section_num, subsection_num)
                            if img_key in image_paths:
                                for img_path in image_paths[img_key]:
                                    md_content.append(("image", img_path))
                            save_markdown_file(
                                md_file,
                                md_content,
                                chapter_num,
                                section_num,
                                subsection_num,
                            )
                            print(
                                f"      ✓ section_{section_num:02d}_{subsection_num:02d}.md"
                            )

    # Create markdown index and CSS
    if ENABLE_MARKDOWN:
        create_markdown_index(chapters, MARKDOWN_DIR)

    # Create index.json for chapter-viewer navigation
    print("\nCreating navigation index...")
    create_navigation_index(chapters, expected_sequence, JSON_DIR)
    print("✓ index.json created")

    print("\n" + "=" * 80)
    print("✓ Book JSON and Markdown generation complete!")
    if ENABLE_MARKDOWN:
        print(f"✓ Markdown files: {MARKDOWN_DIR}/")
    print("=" * 80)


if __name__ == "__main__":
    try:
        build_book_json()
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback

        traceback.print_exc()
        sys.exit(1)
